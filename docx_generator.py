from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import os

SECTION_COLORS = {
    "Price": RGBColor(0, 112, 192),         # Blue
    "People": RGBColor(112, 48, 160),       # Purple
    "Philosophy": RGBColor(0, 176, 80),     # Green
    "Process": RGBColor(255, 192, 0),       # Gold
    "Performance": RGBColor(192, 0, 0),     # Red
}

def clean_markdown(text):
    # Remove leading/trailing "**" or "*" (bold/italic markdown)
    text = re.sub(r'^\*+\s*|\s*\*+$', '', text)
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # Remove bold
    text = re.sub(r'\*(.*?)\*', r'\1', text)      # Remove italic
    return text.strip()

def generate_summary_docx(section_summaries, fund_name="Fund", output_dir="."):
    doc = Document()

    # Title
    doc.add_heading(f"Summary of {fund_name}", 0)

    # Add a table for the summary
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Section'
    hdr_cells[1].text = 'Summary'

    # Style header row
    for i, section in enumerate(['Section', 'Summary']):
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(255, 255, 255)
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Set background color for header (requires python-docx-table-shading, or use default for now)

    for section in ["Price", "People", "Philosophy", "Process", "Performance"]:
        row_cells = table.add_row().cells
        # Section cell
        sec_run = row_cells[0].paragraphs[0].add_run(section)
        sec_run.bold = True
        sec_run.font.size = Pt(11)
        sec_run.font.color.rgb = SECTION_COLORS[section]
        row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Summary cell
        content = section_summaries.get(section, "No information available.")
        content = clean_markdown(content)
        row_cells[1].paragraphs[0].add_run(content).font.size = Pt(11)
        row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Save with fund name in filename
    safe_fund_name = re.sub(r'[^a-zA-Z0-9_\-]', '_', fund_name)
    output_path = os.path.join(output_dir, f"Summary_of_{safe_fund_name}.docx")
    doc.save(output_path)
    return output_path
